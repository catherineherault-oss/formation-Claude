#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Post-traitement du .docx pandoc pour conformité *Décisions Marketing* :
- tout le document en style « Normal » (aucun style de titre prédéfini de Word) ;
- titres présentés manuellement : Titre 1 = TNR 14 gras ; Sous-titre = TNR 12 italique ;
- corps : Times New Roman 12, interligne 1,5, justifié, marges 2,5 cm (gabarit).
Usage : python code/12_format_dm_docx.py <chemin.docx>
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

path = Path(sys.argv[1] if len(sys.argv) > 1 else "outputs/submitted/manuscrit_DM.docx")
d = Document(str(path))

# 1) Style Normal = TNR 12, interligne 1,5, justifié
normal = d.styles["Normal"]
normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing = 1.5; pf.alignment = AL.JUSTIFY

def fmt_runs(p, size, bold=False, italic=False):
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic

TITRE1 = {"Heading 1", "Heading 2", "Title"}   # Titre 1 : TNR 14 gras
SOUS2  = {"Heading 3"}                          # Sous-titre : TNR 12 italique
counts = {"titre1": 0, "sous": 0, "subtitle": 0, "body": 0}

for p in d.paragraphs:
    s = p.style.name
    if s in TITRE1:
        p.style = normal; p.alignment = AL.LEFT; fmt_runs(p, 14, bold=True); counts["titre1"] += 1
    elif s in SOUS2:
        p.style = normal; p.alignment = AL.LEFT; fmt_runs(p, 12, italic=True); counts["sous"] += 1
    elif s == "Subtitle":
        p.style = normal; p.alignment = AL.LEFT; fmt_runs(p, 14, bold=False); counts["subtitle"] += 1
    elif s == "Body Text":
        p.style = normal; counts["body"] += 1     # runs conservés (gras/italique inline préservés)

d.save(str(path))
print(f"Formaté DM : {path.name} — Titre1={counts['titre1']} Sous-titre={counts['sous']} "
      f"Sous-titre1={counts['subtitle']} corps→Normal={counts['body']}")
# contrôle : styles restants
from collections import Counter
c = Counter(p.style.name for p in Document(str(path)).paragraphs)
print("Styles de paragraphes après formatage :", dict(c))
